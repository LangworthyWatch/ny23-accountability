#!/usr/bin/env python3
"""Build a formatted Word doc of the utility rate series captions."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)

doc.styles['Heading 1'].font.size = Pt(20)
doc.styles['Heading 2'].font.size = Pt(15)
doc.styles['Heading 3'].font.size = Pt(12)


def add_meta(label, value):
    p = doc.add_paragraph()
    run_l = p.add_run(f'{label}: ')
    run_l.bold = True
    run_l.font.size = Pt(10)
    run_l.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    run_v = p.add_run(value)
    run_v.font.size = Pt(10)
    run_v.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_caption_block(text):
    """Add the actual caption text in a visually distinct block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.4
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return p


def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('─' * 60)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)


# ──────────────────────────────────────────────────
# TITLE PAGE
# ──────────────────────────────────────────────────

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(80)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Utility Rate Investigation')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Social Media Captions')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run('11-Post Series  |  Facebook & Instagram')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run('langworthywatch.org')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
run.bold = True

doc.add_page_break()

# ──────────────────────────────────────────────────
# POSTING SCHEDULE
# ──────────────────────────────────────────────────

doc.add_heading('Posting Schedule', level=1)

schedule = [
    ('Phase 1: Set the Frame', [
        ('Day 1', 'Post 0', 'What Is Regulatory Capture?', 'post00_regulatory_capture.png'),
        ('Day 2 AM', 'Post 1', 'The Hook — 4.7%', 'post01_hook_4.7pct.png'),
        ('Day 2 PM', 'Post 2', 'The Breakdown', 'post02_breakdown.png'),
    ]),
    ('Phase 2: The Deep Dive', [
        ('Day 3', 'Post 4', '$3.36 vs $21.53/month', 'post04_monthly.png'),
        ('Day 4', 'Post 3', 'CEO Pay vs Your Bill', 'post03_ceo_gap.png'),
        ('Day 5', 'Post 5', 'County Burden', 'post05_burden.png'),
        ('Day 6', 'Post 7', '$75.7M Lobbying', 'post07_lobbying.png'),
    ]),
    ('Phase 3: Accountability', [
        ('Day 7', 'Post 6', 'Complaints +84%', 'post06_complaints.png'),
        ('Day 8', 'Post 8', '18.4% Silence', 'post08_silence.png'),
        ('Day 9 AM', 'Post 9', 'Avangrid Connection', 'post09_connection.png'),
        ('Day 9 PM', 'Post 11', 'Full Fossil Fuel Pipeline', 'post11_pipeline.png'),
        ('Day 10', 'Post 10', 'CTA — Read the Investigation', 'post10_cta.png'),
    ]),
]

for phase_name, posts in schedule:
    doc.add_heading(phase_name, level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Shading Accent 1'
    hdr = table.rows[0].cells
    for i, label in enumerate(['Day', 'Post', 'Title', 'Graphic File']):
        hdr[i].text = label
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for day, post, title, filename in posts:
        row = table.add_row().cells
        row[0].text = day
        row[1].text = post
        row[2].text = title
        row[3].text = filename
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

doc.add_page_break()

# ──────────────────────────────────────────────────
# ALL CAPTIONS
# ──────────────────────────────────────────────────

posts = [
    {
        'num': 'Post 0',
        'title': 'What Is Regulatory Capture?',
        'graphic': 'post00_regulatory_capture.png',
        'day': 'Day 1',
        'hashtags': '#RegulatoryCapture #UtilityRates #NewYork #PSC #FollowTheMoney #DataJournalism',
        'caption': (
            'What is regulatory capture?\n\n'
            'It\'s when the agency meant to regulate an industry is instead controlled by it.\n\n'
            'Here\'s how it works with your electric bill in New York — every step documented with public data:\n\n'
            '1. LOBBY — Energy companies spent $75.7M lobbying Albany. 87.7% on rate-case policy.\n'
            '2. APPOINT — The Governor appoints all 7 PSC commissioners who set your rates.\n'
            '3. REVOLVE — The PSC Chair previously worked at KeySpan Energy (now National Grid).\n'
            '4. APPROVE — The PSC approves rate increases with a guaranteed 9% profit for shareholders.\n'
            '5. PROFIT — Shareholders received $36.8B in dividends. CEO pay grew 64%. Your rates: 42%.\n'
            '6. BLAME — Utilities blame "clean energy" — which is 4.7% of the increase.\n\n'
            'Over the next week, we\'ll show you the data behind every step.\n\n'
            'Sources: COELIG/JCOPE lobbying, PSC commissioner bios, FERC Form 1, EIA\n'
            'langworthywatch.org'
        ),
    },
    {
        'num': 'Post 1',
        'title': 'The Hook — 4.7%',
        'graphic': 'post01_hook_4.7pct.png',
        'day': 'Day 2 (AM)',
        'hashtags': '#NYS #UtilityRates #NYSEG #NY23 #FactCheck #CleanEnergy #DataJournalism',
        'caption': (
            'New York\'s electricity rates are 49.7% above the national average. Politicians keep blaming clean energy mandates.\n\n'
            'We analyzed 8 utility rate case filings totaling $4 BILLION in requested increases.\n\n'
            '-> Clean energy programs: 4.7% ($187M)\n'
            '-> Infrastructure & capital: 33.7% ($1,355M)\n'
            '-> Shareholder profit guarantees: 10.1% ($408M)\n\n'
            'The clean energy surcharge adds about $3.36/month to your bill.\n\n'
            'Full data investigation with downloadable source files at langworthywatch.org'
        ),
    },
    {
        'num': 'Post 2',
        'title': 'The Breakdown',
        'graphic': 'post02_breakdown.png',
        'day': 'Day 2 (PM)',
        'hashtags': '#UtilityBills #NewYork #RateIncrease #PSC #NYSEG #DataDriven',
        'caption': (
            'Your utility filed for a rate increase. Here\'s what they said they need the money for.\n\n'
            'Clean energy programs are the smallest category in the filing.\n\n'
            'Source: PSC rate case filings (dps.ny.gov), 8 cases, 2023-2025\n'
            'Full breakdown + downloadable data: langworthywatch.org'
        ),
    },
    {
        'num': 'Post 3',
        'title': 'CEO Pay vs. Your Bill',
        'graphic': 'post03_ceo_gap.png',
        'day': 'Day 4',
        'hashtags': '#ExecutivePay #UtilityProfits #NationalGrid #NewYork',
        'caption': (
            'Between 2015 and 2024:\n\n'
            '-> CEO compensation at NY\'s top 3 utilities grew 64%\n'
            '-> Your electricity rates grew 42%\n'
            '-> National Grid\'s CEO alone: +78.6%\n\n'
            'Shareholders received $36.8 BILLION in dividends — that\'s 5.3x more than the total clean energy surcharge.\n\n'
            'Source: SEC DEF 14A proxy statements, FERC Form 1\n'
            'langworthywatch.org'
        ),
    },
    {
        'num': 'Post 4',
        'title': '$3.36 vs $21.53/month',
        'graphic': 'post04_monthly.png',
        'day': 'Day 3',
        'hashtags': '#NYSEG #ElectricBill #SouthernTier #NY23',
        'caption': (
            'The System Benefits Charge — the "clean energy surcharge" — adds about $3.36/month to your bill.\n\n'
            'NYSEG is requesting a rate increase that would add $21.53/month.\n\n'
            'If you eliminated the clean energy surcharge entirely, you\'d save $3.36/month.\n'
            'If NYSEG\'s infrastructure request gets approved, you\'d pay $21.53/month more.\n\n'
            'Source: PSC rate case 25-E-0375\n'
            'langworthywatch.org'
        ),
    },
    {
        'num': 'Post 5',
        'title': 'County Burden',
        'graphic': 'post05_burden.png',
        'day': 'Day 5',
        'hashtags': '#SouthernTier #NY23 #NYSEG #RuralNY',
        'caption': (
            'Electricity costs take a bigger bite in lower-income communities.\n\n'
            '-> Lowest income quartile: 2.99% of annual income on electricity\n'
            '-> Highest income quartile: 1.85%\n\n'
            'NYSEG — which serves Steuben, Chemung, Tioga, Allegany, and other Southern Tier counties — '
            'is requesting the steepest increase at 18.4%.\n\n'
            'NYSEG\'s service territory overlaps with the lowest-income counties in NY-23.\n\n'
            'Source: EIA electricity prices, Census ACS 2022\n'
            'langworthywatch.org'
        ),
    },
    {
        'num': 'Post 6',
        'title': 'Complaints +84%',
        'graphic': 'post06_complaints.png',
        'day': 'Day 7',
        'hashtags': '#PSC #ConsumerProtection #UtilityBills',
        'caption': (
            'People are telling regulators they can\'t afford their bills. '
            'The regulators are hearing them — but the utilities keep filing for more.\n\n'
            '-> Statewide complaints: 24,146 to 44,538 (2019-2024)\n'
            '-> Billing complaints grew 120% — 3.2x faster than service complaints\n'
            '-> NYSEG and RG&E: ~95% complaint growth. Same parent company now requesting 18% rate increase.\n\n'
            'Source: PSC consumer complaint data, 2019-2024\n'
            'langworthywatch.org'
        ),
    },
    {
        'num': 'Post 7',
        'title': '$75.7M Lobbying',
        'graphic': 'post07_lobbying.png',
        'day': 'Day 6',
        'hashtags': '#Lobbying #Albany #RateCase #PSC #FollowTheMoney',
        'caption': (
            'New York\'s energy companies spent $75.7 million lobbying in Albany over the past decade.\n\n'
            '-> 87.7% of that ($62.9M) was specifically about rate cases and energy policy\n'
            '-> Annual lobbying spend nearly tripled from 2011 to 2024\n'
            '-> Top spenders: Business Council $8.9M, Entergy $5.6M, National Grid $5.5M\n\n'
            'All 7 PSC commissioners who approve rate increases are appointed by the Governor.\n\n'
            'Source: COELIG/JCOPE lobbying disclosures\n'
            'langworthywatch.org'
        ),
    },
    {
        'num': 'Post 8',
        'title': '18.4% Silence',
        'graphic': 'post08_silence.png',
        'day': 'Day 8',
        'hashtags': '#NY23 #NYSEG #RateHike #Accountability',
        'caption': (
            'NYSEG — the utility serving most of NY-23 — is requesting an 18.4% rate increase.\n\n'
            'That\'s the steepest pending increase of any major utility in the state.\n\n'
            'Rep. Langworthy has repeatedly blamed clean energy mandates for high bills. '
            'He has not publicly addressed the NYSEG rate case currently before the PSC.\n\n'
            'Of NYSEG\'s rate increase request, clean energy programs account for 4.7%. '
            'Infrastructure and capital investment: 33.7%.\n\n'
            'langworthywatch.org'
        ),
    },
    {
        'num': 'Post 9',
        'title': 'Avangrid Connection',
        'graphic': 'post09_connection.png',
        'day': 'Day 9 (AM)',
        'hashtags': '#FollowTheMoney #NY23 #NYSEG #CampaignFinance',
        'caption': (
            'Avangrid — the parent company of NYSEG — donated to Langworthy\'s congressional campaign.\n\n'
            'Avangrid is currently requesting an 18.4% rate increase that would hit '
            'NY-23\'s lowest-income counties the hardest.\n\n'
            'Langworthy has not publicly commented on the pending rate case.\n\n'
            'Source: FEC individual contributions, C00817932\n'
            'langworthywatch.org'
        ),
    },
    {
        'num': 'Post 11',
        'title': 'The Full Fossil Fuel Pipeline',
        'graphic': 'post11_pipeline.png',
        'day': 'Day 9 (PM)',
        'hashtags': '#FollowTheMoney #NY23 #RegulatoryCapture #UtilityRates #CampaignFinance #PSC',
        'caption': (
            'Post 9 showed Avangrid\'s $1,500 donation. Here\'s the full picture.\n\n'
            'In the 2024 cycle, Langworthy received $66,466 from the oil and gas industry — '
            'his 8th largest donor sector. $50,000 from PACs. $16,466 from individuals.\n\n'
            'Contributors include Koch Industries PAC, Marathon Petroleum PAC, and United Refining Co. ($21,900).\n\n'
            'Langworthy has appeared alongside Republican colleagues lobbying against clean energy mandates, '
            'calling them a primary driver of rising electricity bills. His solution: the Energy Choice Act '
            '(H.R. 3699), framed as restoring consumer choice and protecting natural gas access.\n\n'
            'That bill was drafted by NEFI, the National Energy & Fuels Institute — a fossil fuel trade group. '
            '30+ fossil fuel companies endorsed it. His 2024 environmental score from the League of Conservation Voters: 0%.\n\n'
            'NYSEG — whose parent company Avangrid donated to his campaign — is currently asking the PSC '
            'for an 18.4% rate increase affecting NY-23 customers. Langworthy has made no public statement on the case.\n\n'
            'Sources: FEC (fec.gov), NYSBOE, OpenSecrets, LCV Scorecard, COELIG/JCOPE\n'
            'Full investigation: langworthywatch.org'
        ),
    },
    {
        'num': 'Post 10',
        'title': 'Read the Full Investigation',
        'graphic': 'post10_cta.png',
        'day': 'Day 10',
        'hashtags': '#LangworthyWatch #DataJournalism #Transparency #NY23',
        'caption': (
            'We published the full data investigation with:\n\n'
            '-> 6 interactive charts\n'
            '-> 4 downloadable CSV datasets\n'
            '-> Every source linked\n'
            '-> No paywalls, no ads, no party affiliation\n\n'
            'langworthywatch.org/fact-checks/2026-03-14-nys-utility-rates-data-investigation/\n\n'
            'Share if you think your neighbors should see this data.'
        ),
    },
]

FULL_URL = 'https://langworthywatch.org/fact-checks/2026-03-14-nys-utility-rates-data-investigation'

doc.add_heading('Captions — Ready to Copy & Paste', level=1)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
run = p.add_run('Each caption below includes the full investigation URL and hashtags. '
                 'Select the entire shaded block and paste directly into Facebook or Instagram.')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

for post in posts:
    doc.add_heading(f'{post["num"]}: {post["title"]}', level=2)

    add_meta('Graphic', post['graphic'])
    add_meta('Schedule', post['day'])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run('COPY BELOW:')
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    # Build the full copy-paste block: caption + URL + hashtags
    caption = post['caption']
    # Replace bare langworthywatch.org links with full URL where appropriate
    if 'Full investigation: langworthywatch.org' in caption:
        caption = caption.replace('Full investigation: langworthywatch.org',
                                  f'Full investigation: {FULL_URL}')
    elif 'Full data investigation with downloadable source files at langworthywatch.org' in caption:
        caption = caption.replace(
            'Full data investigation with downloadable source files at langworthywatch.org',
            f'Full data investigation with downloadable source files:\n{FULL_URL}')
    elif caption.rstrip().endswith('langworthywatch.org'):
        caption = caption.rstrip().rsplit('langworthywatch.org', 1)[0] + FULL_URL

    full_block = caption + '\n\n' + post['hashtags']

    add_caption_block(full_block)
    add_divider()

# ──────────────────────────────────────────────────
# NOTES PAGE
# ──────────────────────────────────────────────────

doc.add_page_break()
doc.add_heading('Platform Notes', level=1)

notes = [
    ('Instagram', 'Upload as regular posts (not Stories). 1080x1350 portrait is feed-optimized. '
     'Stories disappear in 24h — use them only for re-shares after the series runs.'),
    ('Facebook', 'Image posts upload natively. The 1080x1350 portrait format works well in feed. '
     'For the GIF/video post (if applicable), convert to MP4 first — Business Suite does not accept GIFs.'),
    ('If challenged', 'Point to langworthywatch.org for full sourced investigation with downloadable data. '
     'All numbers come from PSC rate filings, FERC Form 1, SEC proxy statements, EIA, and COELIG/JCOPE — '
     'all public records.'),
    ('Tone', 'Present data. Let readers draw conclusions. No zingers, no editorial framing. '
     'The numbers are strong enough on their own.'),
]

for label, text in notes:
    p = doc.add_paragraph()
    run_l = p.add_run(f'{label}: ')
    run_l.bold = True
    run_l.font.size = Pt(10)
    run_v = p.add_run(text)
    run_v.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)

# Save
out_path = '/Users/zachbeaudoin/projects/Langworthywatch/langworthy-tracker/social-media/utility-rate-series/Utility_Rate_Series_Captions.docx'
doc.save(out_path)
size_kb = os.path.getsize(out_path) // 1024
print(f'Saved: {out_path} ({size_kb} KB)')
